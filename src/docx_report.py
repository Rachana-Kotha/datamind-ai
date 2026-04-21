"""
DataMind AI — Word (.docx) Report Generator
Creates a professional branded Word document using python-docx.
"""

import io
from datetime import datetime
from typing import List, Dict, Optional

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re


# ── Brand colors (RGB) ────────────────────────────────────────────────────────
PURPLE      = RGBColor(0x4F, 0x46, 0xE5)
PURPLE_DARK = RGBColor(0x37, 0x30, 0xA3)
PURPLE_LITE = RGBColor(0xEE, 0xF2, 0xFF)
TEAL        = RGBColor(0x0F, 0x6E, 0x56)
AMBER       = RGBColor(0xB4, 0x53, 0x09)
RED         = RGBColor(0xB9, 0x1C, 0x1C)
GREEN       = RGBColor(0x15, 0x80, 0x3D)
GRAY_DARK   = RGBColor(0x1F, 0x29, 0x37)
GRAY_MID    = RGBColor(0x6B, 0x72, 0x80)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)


def _set_cell_bg(cell, hex_color: str):
    """Set cell background color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell_border(cell, edges=("top", "bottom", "left", "right"),
                     color="E5E7EB", size="4"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in edges:
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _heading(doc, text: str, level: int = 1,
             color: RGBColor = PURPLE_DARK):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = color
    return h


def _para(doc, text: str, bold: bool = False, italic: bool = False,
          color: RGBColor = GRAY_DARK, size: int = 10, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    # Handle **bold** markdown inline
    parts = re.split(r"\*\*(.+?)\*\*", text)
    for i, part in enumerate(parts):
        if part == "":
            continue
        run = p.add_run(part)
        run.bold = bold or (i % 2 == 1)
        run.italic = italic
        run.font.color.rgb = color
        run.font.size = Pt(size)
    return p


def _hr(doc, color: str = "4F46E5"):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def generate_docx(
    dataset_name: str,
    target_col: str,
    task_type: str,
    narrative: str,
    kpis: List[Dict],
    ml_findings: Dict,
    eda_findings: Dict,
    insights: List[Dict],
    risks: List[Dict],
    debate: List[Dict],
    share_url: Optional[str] = None,
    output_path: Optional[str] = None,
) -> bytes:
    """Build the Word report and return as bytes."""
    doc = Document()
    now = datetime.now().strftime("%B %d, %Y  %H:%M")

    # ── Page setup ────────────────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width  = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = section.right_margin = Inches(0.9)
    section.top_margin  = section.bottom_margin = Inches(0.85)

    # ── Cover ─────────────────────────────────────────────────────────────────
    doc.add_paragraph()
    t_para = doc.add_paragraph()
    t_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t_para.add_run("DataMind AI")
    run.bold = True
    run.font.size = Pt(36)
    run.font.color.rgb = PURPLE

    sub = doc.add_paragraph("Data Intelligence Report")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in sub.runs:
        r.font.size = Pt(16)
        r.font.color.rgb = GRAY_MID

    _hr(doc)

    shape = eda_findings.get("shape", {})
    meta_rows = [
        ("Dataset", dataset_name),
        ("Target variable", target_col),
        ("Task type", task_type.title()),
        ("Generated", now),
        ("Best model", ml_findings.get("best_model", "N/A")),
        ("Best score", f"{ml_findings.get('best_score', 0):.4f} ({ml_findings.get('metric_name', 'score')})"),
    ]
    if share_url:
        meta_rows.append(("Report URL", share_url))

    meta_table = doc.add_table(rows=len(meta_rows), cols=2)
    meta_table.style = "Table Grid"
    for i, (key, val) in enumerate(meta_rows):
        cells = meta_table.rows[i].cells
        cells[0].text = key
        cells[0].paragraphs[0].runs[0].bold = True
        cells[0].paragraphs[0].runs[0].font.color.rgb = PURPLE_DARK
        cells[1].text = val
        cells[1].paragraphs[0].runs[0].font.color.rgb = GRAY_DARK
        _set_cell_border(cells[0])
        _set_cell_border(cells[1])

    doc.add_paragraph()

    # Summary stats row
    summary_table = doc.add_table(rows=1, cols=5)
    labels = [
        (f"{shape.get('rows', 0):,}", "Rows"),
        (str(shape.get("cols", 0)), "Columns"),
        (str(len(kpis)), "Top KPIs"),
        (str(len(insights)), "Insights"),
        (str(len(risks)), "Risks"),
    ]
    for i, (val, lbl) in enumerate(labels):
        cell = summary_table.rows[0].cells[i]
        _set_cell_bg(cell, "EEF2FF")
        _set_cell_border(cell, color="4F46E5")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run(f"\n{val}\n")
        r1.bold = True
        r1.font.size = Pt(18)
        r1.font.color.rgb = PURPLE if i < 4 else RED if risks else GREEN
        r2 = p.add_run(lbl)
        r2.font.size = Pt(9)
        r2.font.color.rgb = GRAY_MID

    doc.add_page_break()

    # ── Data narrative ────────────────────────────────────────────────────────
    _heading(doc, "The Data Story", level=1)
    _hr(doc)
    for para_text in narrative.split("\n\n"):
        if para_text.strip():
            _para(doc, para_text.strip())

    doc.add_paragraph()

    # ── KPIs ─────────────────────────────────────────────────────────────────
    _heading(doc, "Top 5 KPIs", level=1)
    _hr(doc)

    for kpi in kpis:
        col_name = kpi.get("title", kpi.get("column", "KPI"))
        mean_val = kpi.get("mean", 0)
        trend = kpi.get("trend_direction", "flat")
        trend_pct = kpi.get("trend_pct", 0)
        corr = kpi.get("target_correlation")

        if abs(mean_val) >= 1_000_000:
            val_str = f"{mean_val/1_000_000:.2f}M"
        elif abs(mean_val) >= 1_000:
            val_str = f"{mean_val/1_000:.2f}K"
        else:
            val_str = f"{mean_val:.3f}"

        kpi_table = doc.add_table(rows=2, cols=4)
        kpi_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header row
        hrow = kpi_table.rows[0]
        _set_cell_bg(hrow.cells[0], "4F46E5")
        for cell in hrow.cells:
            _set_cell_bg(cell, "4F46E5")
            _set_cell_border(cell, color="3730A3")

        hrow.cells[0].merge(hrow.cells[3])
        hp = hrow.cells[0].paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        hr = hp.add_run(f"  {col_name}")
        hr.bold = True
        hr.font.color.rgb = WHITE
        hr.font.size = Pt(13)

        # Data row
        drow = kpi_table.rows[1]
        stats = [
            (val_str, "Average"),
            (f"{'▲' if trend == 'up' else '▼' if trend == 'down' else '●'} {trend_pct:+.1f}%", "Trend"),
            (f"r={corr:.2f}" if corr else "—", "vs Target"),
            (f"{kpi.get('count', 0):,}", "Count"),
        ]
        for i, (v, l) in enumerate(stats):
            _set_cell_bg(drow.cells[i], "EEF2FF")
            _set_cell_border(drow.cells[i], color="4F46E5")
            cp = drow.cells[i].paragraphs[0]
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cr1 = cp.add_run(f"\n{v}\n")
            cr1.bold = True
            cr1.font.size = Pt(16)
            cr1.font.color.rgb = PURPLE
            cr2 = cp.add_run(l)
            cr2.font.size = Pt(8)
            cr2.font.color.rgb = GRAY_MID

        doc.add_paragraph()
        _heading(doc, f"● {col_name}", level=3, color=PURPLE_DARK)
        _para(doc, kpi.get("story", ""))
        doc.add_paragraph()

    doc.add_page_break()

    # ── Model leaderboard ────────────────────────────────────────────────────
    _heading(doc, "Model Leaderboard", level=1)
    _hr(doc)

    leaderboard = ml_findings.get("leaderboard", [])
    if leaderboard:
        if task_type == "classification":
            headers = ["Rank", "Model", "Accuracy", "F1", "CV Mean", "CV Std"]
            rows = [[str(i+1), r.get("model",""), f"{r.get('accuracy',0):.4f}",
                     f"{r.get('f1',0):.4f}", f"{r.get('cv_mean',0):.4f}", f"±{r.get('cv_std',0):.4f}"]
                    for i, r in enumerate(leaderboard)]
        else:
            headers = ["Rank", "Model", "R²", "RMSE", "CV Mean", "CV Std"]
            rows = [[str(i+1), r.get("model",""), f"{r.get('r2',0):.4f}",
                     f"{r.get('rmse',0):.4f}", f"{r.get('cv_mean',0):.4f}", f"±{r.get('cv_std',0):.4f}"]
                    for i, r in enumerate(leaderboard)]

        lb_table = doc.add_table(rows=1+len(rows), cols=len(headers))
        lb_table.style = "Table Grid"
        for j, h in enumerate(headers):
            cell = lb_table.rows[0].cells[j]
            _set_cell_bg(cell, "4F46E5")
            cell.paragraphs[0].add_run(h).font.color.rgb = WHITE
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(9)

        for i, row in enumerate(rows):
            bg = "EEF2FF" if i == 0 else ("F9FAFB" if i % 2 == 0 else "FFFFFF")
            for j, val in enumerate(row):
                cell = lb_table.rows[i+1].cells[j]
                _set_cell_bg(cell, bg)
                run = cell.paragraphs[0].add_run(val)
                run.font.size = Pt(9)
                if i == 0:
                    run.bold = True
                    run.font.color.rgb = PURPLE_DARK

    doc.add_paragraph()

    # ── Insights ─────────────────────────────────────────────────────────────
    _heading(doc, "Key Insights", level=1)
    _hr(doc)
    for ins in insights:
        t = ins.get("type", "insight")
        color = {"warning": AMBER, "success": GREEN, "insight": TEAL}.get(t, TEAL)
        icon = {"warning": "⚠ ", "success": "✓ ", "insight": "💡 "}.get(t, "• ")
        ih = doc.add_paragraph()
        r = ih.add_run(f"{icon}{ins.get('title', '')}")
        r.bold = True
        r.font.color.rgb = color
        r.font.size = Pt(11)
        _para(doc, ins.get("detail", ""), color=GRAY_DARK)
        doc.add_paragraph()

    # ── Risks ─────────────────────────────────────────────────────────────────
    if risks:
        _heading(doc, "Risk Audit", level=2, color=RED)
        for risk in risks:
            sev = risk.get("severity", "low")
            color = RED if sev in ["high", "critical"] else AMBER if sev == "medium" else GREEN
            rh = doc.add_paragraph()
            r = rh.add_run(f"[{sev.upper()}] {risk.get('risk', '')}")
            r.bold = True
            r.font.color.rgb = color
            r.font.size = Pt(10)
            _para(doc, risk.get("detail", ""), color=GRAY_DARK)

    # ── Share URL ─────────────────────────────────────────────────────────────
    if share_url:
        _hr(doc)
        _heading(doc, "Share This Report", level=2)
        p = doc.add_paragraph()
        r = p.add_run(f"Report URL: {share_url}")
        r.font.color.rgb = PURPLE
        r.font.size = Pt(10)

    # ── Save ─────────────────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    docx_bytes = buf.getvalue()
    if output_path:
        with open(output_path, "wb") as f:
            f.write(docx_bytes)
    return docx_bytes
